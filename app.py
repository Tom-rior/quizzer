import streamlit as st
import pandas as pd
import random
from docx import Document
from docx.shared import Pt
import io

# Page Title
st.title("📝 Quiz Maker Pro")
st.write("Upload a CSV and download a ready-to-use Quiz Word doc!")

# 1. File Upload
uploaded_files = st.file_uploader("Choose CSV file(s)", type="csv", accept_multiple_files=True)

if uploaded_files:
    dfs = []
    for uploaded_file in uploaded_files:
        df = pd.read_csv(uploaded_file)
        
        # Rename columns to match your logic
        rename_map = {
            df.columns[0]: 'Question',
            df.columns[1]: 'A',
            df.columns[2]: 'B',
            df.columns[3]: 'C',
            df.columns[4]: 'D',
            df.columns[5]: 'Answer'
        }
        df = df.rename(columns=rename_map)
        dfs.append(df)

    all_df = pd.concat(dfs, ignore_index=True).drop_duplicates(subset=['Question']).dropna()
    st.success(f"Loaded {len(all_df)} unique questions!")

    # 2. Quiz Settings
    num_questions = st.number_input("How many questions?", min_value=1, max_value=len(all_df), value=min(10, len(all_df)))
    neg_marking = st.checkbox("Enable negative marking (-1/3)?")

    if st.button("Generate Quiz"):
        # Random selection
        selected = all_df.sample(n=num_questions).to_dict('records')

        # 3. Create DOCX in memory
        doc = Document()
        doc.add_heading('Quiz', 0)
        if neg_marking:
            doc.add_paragraph('Marking: +1 correct, -1/3 incorrect, 0 unattempted')

        # Add Questions
        for i, q in enumerate(selected, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i}. {q['Question']}")
            run.bold = True
            doc.add_paragraph(f"A. {q['A']}")
            doc.add_paragraph(f"B. {q['B']}")
            doc.add_paragraph(f"C. {q['C']}")
            doc.add_paragraph(f"D. {q['D']}")
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Add Answer Key
        doc.add_page_break()
        doc.add_heading('Answer Key', level=1)
        for i, q in enumerate(selected, 1):
            doc.add_paragraph(f"Q{i}. {q['Answer']}")

        # Save to a buffer (so user can download)
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label="📥 Download Quiz (DOCX)",
            data=bio.getvalue(),
            file_name="quiz.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
